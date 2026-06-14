"""Document exporters for the BTP OpenClaw Cockpit.

Converts a ``Document`` ORM row plus an optional ``CompanySettings`` row into
binary bytes in one of three formats:

* **PDF**  — reportlab (pure-Python, no system binary required).
* **DOCX** — python-docx.
* **XLSX** — openpyxl  (particularly useful for the quote / devis type).

Public API
----------
::

    bytes_, media_type, filename = export_document(doc, company, fmt)

``fmt`` is case-insensitive and must be one of ``"pdf"``, ``"docx"``, ``"xlsx"``.

Content schemas handled
-----------------------
Each document type maps to a known JSON ``content`` shape produced by the
corresponding agent.  The exporter is defensive: any missing key is silently
skipped so a partial or stub ``content`` dict never causes a crash.

``quote``
    content["lines"]      list of {label, qty, unit, unit_price_ht, total_ht}
    content["total_ht"]   float
    content["tva_rate"]   float (default 0.20)
    content["total_tva"]  float
    content["total_ttc"]  float
    content["hypotheses"] list[str]

``site_report``
    content["date"]       str
    content["present"]    list[str]
    content["constats"]   list[str]
    content["actions"]    list[str]
    content["reserves"]   list[str]

``tender_response``
    content["pieces_demandees"] list[str]
    content["criteres"]         list[str]
    content["delais"]           str
    content["points_vigilance"] list[str]

``photo_report``
    content["observations"]      str
    content["travaux_visibles"]  list[str]
    content["points_attention"]  list[str]

Any other ``document_type``
    Generic rendering: each top-level key becomes a section heading; str values
    are rendered as a paragraph, list values as a bullet list.
"""

from __future__ import annotations

import io
import re
from datetime import datetime
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from models import CompanySettings, Document


# ---------------------------------------------------------------------------
# Public exception
# ---------------------------------------------------------------------------


class ExportUnavailable(Exception):
    """Raised when a required export library is not installed."""


# ---------------------------------------------------------------------------
# Small helpers
# ---------------------------------------------------------------------------


def _s(value: Any, default: str = "") -> str:
    """Coerce any value to a stripped string, never None."""
    if value is None:
        return default
    return str(value).strip()


def _f(value: Any, default: float = 0.0) -> float:
    """Best-effort float coercion."""
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _lst(value: Any) -> list:
    """Return value if it is a list, else an empty list."""
    return value if isinstance(value, list) else []


def _eur(amount: float) -> str:
    """Format a float as a Euro amount with 2 decimal places."""
    return f"{amount:,.2f} €".replace(",", " ")


def _today() -> str:
    return datetime.now().strftime("%d/%m/%Y")


def _filename_stem(title: str) -> str:
    """Derive a safe ASCII filename stem from a document title."""
    slug = re.sub(r"[^\w-]", "_", title, flags=re.ASCII)
    slug = re.sub(r"_+", "_", slug).strip("_")
    return slug[:60] or "document"


# ---------------------------------------------------------------------------
# Company header extraction
# ---------------------------------------------------------------------------


def _ch(company: Any) -> dict[str, Any]:
    """Return a plain dict of display fields from a CompanySettings row.

    Works with the ORM model, any object with matching attributes, or ``None``.
    """
    if company is None:
        return {
            "name": "Mon Entreprise BTP",
            "address": "",
            "email": "",
            "phone": "",
            "siret": "",
            "vat_number": "",
            "legal_mentions": "",
            "default_tva_rate": 0.20,
        }
    return {
        "name": _s(getattr(company, "company_name", "Mon Entreprise BTP")) or "Mon Entreprise BTP",
        "address": _s(getattr(company, "address", None)),
        "email": _s(getattr(company, "email", None)),
        "phone": _s(getattr(company, "phone", None)),
        "siret": _s(getattr(company, "siret", None)),
        "vat_number": _s(getattr(company, "vat_number", None)),
        "legal_mentions": _s(getattr(company, "legal_mentions", None)),
        "default_tva_rate": _f(getattr(company, "default_tva_rate", 0.20), 0.20),
    }


def _company_lines(company: Any) -> list[str]:
    """Return non-empty display lines for the company block."""
    c = _ch(company)
    lines: list[str] = [c["name"]]
    for field in ("address", "phone", "email"):
        if c[field]:
            prefix = {"phone": "Tél : ", "email": "Email : "}.get(field, "")
            lines.append(f"{prefix}{c[field]}")
    if c["siret"]:
        lines.append(f"SIRET : {c['siret']}")
    if c["vat_number"]:
        lines.append(f"N° TVA : {c['vat_number']}")
    return lines


# ---------------------------------------------------------------------------
# Quote line parsing (shared across formats)
# ---------------------------------------------------------------------------


def _parse_quote(content: dict, default_tva: float) -> tuple[list[dict], float, float, float, float]:
    """Extract quote lines and totals from a content dict.

    Returns (lines, tva_rate, total_ht, total_tva, total_ttc).

    Each item in *lines* is guaranteed to have: label, qty, unit,
    unit_price_ht, total_ht.
    """
    raw = _lst(content.get("lines"))
    tva_rate = _f(content.get("tva_rate"), default_tva)

    lines: list[dict] = []
    for row in raw:
        if not isinstance(row, dict):
            continue
        qty = _f(row.get("qty"))
        unit_price_ht = _f(row.get("unit_price_ht"))
        lines.append(
            {
                "label": _s(row.get("label")) or "Ligne sans libellé",
                "qty": qty,
                "unit": _s(row.get("unit")) or "u",
                "unit_price_ht": unit_price_ht,
                "total_ht": _f(row.get("total_ht"), qty * unit_price_ht),
            }
        )

    total_ht = _f(content.get("total_ht"), sum(ln["total_ht"] for ln in lines))
    total_tva = _f(content.get("total_tva"), total_ht * tva_rate)
    total_ttc = _f(content.get("total_ttc"), total_ht + total_tva)
    return lines, tva_rate, total_ht, total_tva, total_ttc


# ===========================================================================
# PDF exporter — reportlab Platypus
# ===========================================================================


def _export_pdf(doc: Any, company: Any) -> bytes:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
    except ImportError as exc:
        raise ExportUnavailable(
            "reportlab n'est pas installé (pip install reportlab)."
        ) from exc

    AMBER = colors.HexColor("#d97706")
    DARK = colors.HexColor("#1a1a1a")
    GREY_ROW = colors.HexColor("#f3f4f6")

    buf = io.BytesIO()
    page_w, _ = A4
    usable_w = page_w - 4 * cm  # 2 cm margin each side

    pdf = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    ss = getSampleStyleSheet()

    st_company = ParagraphStyle("company", parent=ss["Normal"], fontSize=10, leading=14, textColor=DARK)
    st_h1 = ParagraphStyle("h1", parent=ss["Heading1"], fontSize=14, textColor=DARK, spaceAfter=4)
    st_h2 = ParagraphStyle("h2", parent=ss["Heading2"], fontSize=11, textColor=colors.HexColor("#374151"), spaceAfter=4)
    st_body = ParagraphStyle("body", parent=ss["Normal"], fontSize=9, leading=13)
    st_legal = ParagraphStyle("legal", parent=ss["Normal"], fontSize=7, leading=10, textColor=colors.HexColor("#6b7280"))

    def bullet(text: str) -> Paragraph:
        return Paragraph(f"• {text}", st_body)

    content: dict = getattr(doc, "content", None) or {}
    doc_type: str = _s(getattr(doc, "document_type", ""))
    title: str = _s(getattr(doc, "title", "Document")) or "Document"
    c = _ch(company)

    els: list = []

    # --- Company block ---
    for line in _company_lines(company):
        els.append(Paragraph(line, st_company))
    els.append(Spacer(1, 0.4 * cm))

    # --- Document title + date ---
    els.append(Paragraph(title, st_h1))
    els.append(Paragraph(f"Date : {_today()}", st_body))
    els.append(Spacer(1, 0.5 * cm))

    # --- Body by type ---
    if doc_type == "quote":
        lines, tva_rate, total_ht, total_tva, total_ttc = _parse_quote(content, c["default_tva_rate"])

        els.append(Paragraph("Détail des prestations", st_h2))

        # Line-items table
        hdr_row = ["Désignation", "Qté", "Unité", "PU HT", "Total HT"]
        col_w = [usable_w - 5 * cm, 1.2 * cm, 1.4 * cm, 2.2 * cm, 2.2 * cm]
        rows: list[list] = [hdr_row] + [
            [
                ln["label"],
                f"{ln['qty']:g}",
                ln["unit"],
                _eur(ln["unit_price_ht"]),
                _eur(ln["total_ht"]),
            ]
            for ln in lines
        ]
        tbl = Table(rows, colWidths=col_w, repeatRows=1)
        tbl.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), AMBER),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
                    ("ALIGN", (0, 0), (0, -1), "LEFT"),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#d1d5db")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, GREY_ROW]),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        els.append(tbl)
        els.append(Spacer(1, 0.35 * cm))

        # Totals
        tva_pct = f"{tva_rate * 100:.0f} %"
        totals_rows = [
            ["Total HT", _eur(total_ht)],
            [f"TVA ({tva_pct})", _eur(total_tva)],
            ["Total TTC", _eur(total_ttc)],
        ]
        tot_tbl = Table(totals_rows, colWidths=[3.5 * cm, 3 * cm], hAlign="RIGHT")
        tot_tbl.setStyle(
            TableStyle(
                [
                    ("ALIGN", (0, 0), (-1, -1), "RIGHT"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("FONTNAME", (0, 2), (-1, 2), "Helvetica-Bold"),
                    ("LINEABOVE", (0, 2), (-1, 2), 0.75, AMBER),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]
            )
        )
        els.append(tot_tbl)

        hypotheses = _lst(content.get("hypotheses"))
        if hypotheses:
            els.append(Spacer(1, 0.5 * cm))
            els.append(Paragraph("Hypothèses", st_h2))
            for h in hypotheses:
                els.append(bullet(_s(h)))

    elif doc_type == "site_report":
        if date := _s(content.get("date")):
            els.append(Paragraph(f"Date de visite : {date}", st_body))
            els.append(Spacer(1, 0.3 * cm))
        for section_title, key in [
            ("Participants", "present"),
            ("Constats", "constats"),
            ("Actions à mener", "actions"),
            ("Réserves", "reserves"),
        ]:
            items = _lst(content.get(key))
            if items:
                els.append(Paragraph(section_title, st_h2))
                for item in items:
                    els.append(bullet(_s(item)))
                els.append(Spacer(1, 0.2 * cm))

    elif doc_type == "tender_response":
        for section_title, key in [
            ("Pièces demandées", "pieces_demandees"),
            ("Critères de sélection", "criteres"),
            ("Points de vigilance", "points_vigilance"),
        ]:
            items = _lst(content.get(key))
            if items:
                els.append(Paragraph(section_title, st_h2))
                for item in items:
                    els.append(bullet(_s(item)))
                els.append(Spacer(1, 0.2 * cm))
        if delais := _s(content.get("delais")):
            els.append(Paragraph("Délais", st_h2))
            els.append(Paragraph(delais, st_body))

    elif doc_type == "photo_report":
        if obs := _s(content.get("observations")):
            els.append(Paragraph("Observations", st_h2))
            els.append(Paragraph(obs, st_body))
            els.append(Spacer(1, 0.2 * cm))
        for section_title, key in [
            ("Travaux visibles", "travaux_visibles"),
            ("Points d'attention", "points_attention"),
        ]:
            items = _lst(content.get(key))
            if items:
                els.append(Paragraph(section_title, st_h2))
                for item in items:
                    els.append(bullet(_s(item)))
                els.append(Spacer(1, 0.2 * cm))

    else:
        # Generic: any top-level key → section
        for key, value in content.items():
            section_label = key.replace("_", " ").capitalize()
            els.append(Paragraph(section_label, st_h2))
            if isinstance(value, list):
                for item in value:
                    els.append(bullet(_s(item)))
            elif isinstance(value, dict):
                for k, v in value.items():
                    els.append(Paragraph(f"<b>{k} :</b> {v}", st_body))
            else:
                els.append(Paragraph(_s(value), st_body))
            els.append(Spacer(1, 0.25 * cm))

    # --- Legal mentions ---
    if legal := c["legal_mentions"]:
        els.append(Spacer(1, 0.6 * cm))
        els.append(Paragraph("Mentions légales", st_h2))
        els.append(Paragraph(legal.replace("\n", "<br/>"), st_legal))

    pdf.build(els)
    buf.seek(0)
    return buf.read()


# ===========================================================================
# DOCX exporter — python-docx
# ===========================================================================


def _export_docx(doc: Any, company: Any) -> bytes:
    try:
        from docx import Document as DocxDoc
        from docx.shared import Pt, RGBColor
    except ImportError as exc:
        raise ExportUnavailable(
            "python-docx n'est pas installé (pip install python-docx)."
        ) from exc

    c = _ch(company)
    content: dict = getattr(doc, "content", None) or {}
    doc_type: str = _s(getattr(doc, "document_type", ""))
    title: str = _s(getattr(doc, "title", "Document")) or "Document"

    d = DocxDoc()

    # Page margins (in EMU: 1 cm = 360 000 EMU)
    for section in d.sections:
        section.top_margin = 720000      # 2 cm
        section.bottom_margin = 720000
        section.left_margin = 900000     # 2.5 cm
        section.right_margin = 900000

    # --- Company block ---
    for line in _company_lines(company):
        p = d.add_paragraph(line)
        if p.runs:
            p.runs[0].font.size = Pt(10)

    d.add_paragraph()

    # --- Document title ---
    h1 = d.add_heading(title, level=1)
    if h1.runs:
        h1.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    d.add_paragraph(f"Date : {_today()}")
    d.add_paragraph()

    def add_h2(text: str) -> None:
        h = d.add_heading(text, level=2)
        if h.runs:
            h.runs[0].font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    def add_bullets(items: list) -> None:
        for item in items:
            d.add_paragraph(_s(item), style="List Bullet")

    # --- Body by type ---
    if doc_type == "quote":
        lines, tva_rate, total_ht, total_tva, total_ttc = _parse_quote(content, c["default_tva_rate"])

        add_h2("Détail des prestations")

        if lines:
            tbl = d.add_table(rows=1 + len(lines), cols=5)
            tbl.style = "Table Grid"
            for i, hdr in enumerate(["Désignation", "Qté", "Unité", "PU HT", "Total HT"]):
                cell = tbl.rows[0].cells[i]
                cell.text = hdr
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
            for row_i, ln in enumerate(lines, start=1):
                cells = tbl.rows[row_i].cells
                cells[0].text = ln["label"]
                cells[1].text = f"{ln['qty']:g}"
                cells[2].text = ln["unit"]
                cells[3].text = _eur(ln["unit_price_ht"])
                cells[4].text = _eur(ln["total_ht"])
                for cell in cells:
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(9)

        d.add_paragraph()
        tva_pct = f"{tva_rate * 100:.0f} %"
        for label, amount in [
            ("Total HT", total_ht),
            (f"TVA ({tva_pct})", total_tva),
            ("Total TTC", total_ttc),
        ]:
            p = d.add_paragraph()
            run = p.add_run(f"{label} : {_eur(amount)}")
            if label.startswith("Total TTC"):
                run.bold = True

        hypotheses = _lst(content.get("hypotheses"))
        if hypotheses:
            d.add_paragraph()
            add_h2("Hypothèses")
            add_bullets(hypotheses)

    elif doc_type == "site_report":
        if date := _s(content.get("date")):
            d.add_paragraph(f"Date de visite : {date}")
        for section_title, key in [
            ("Participants", "present"),
            ("Constats", "constats"),
            ("Actions à mener", "actions"),
            ("Réserves", "reserves"),
        ]:
            items = _lst(content.get(key))
            if items:
                d.add_paragraph()
                add_h2(section_title)
                add_bullets(items)

    elif doc_type == "tender_response":
        for section_title, key in [
            ("Pièces demandées", "pieces_demandees"),
            ("Critères de sélection", "criteres"),
            ("Points de vigilance", "points_vigilance"),
        ]:
            items = _lst(content.get(key))
            if items:
                d.add_paragraph()
                add_h2(section_title)
                add_bullets(items)
        if delais := _s(content.get("delais")):
            d.add_paragraph()
            add_h2("Délais")
            d.add_paragraph(delais)

    elif doc_type == "photo_report":
        if obs := _s(content.get("observations")):
            add_h2("Observations")
            d.add_paragraph(obs)
        for section_title, key in [
            ("Travaux visibles", "travaux_visibles"),
            ("Points d'attention", "points_attention"),
        ]:
            items = _lst(content.get(key))
            if items:
                d.add_paragraph()
                add_h2(section_title)
                add_bullets(items)

    else:
        for key, value in content.items():
            section_label = key.replace("_", " ").capitalize()
            d.add_paragraph()
            add_h2(section_label)
            if isinstance(value, list):
                add_bullets(value)
            elif isinstance(value, dict):
                for k, v in value.items():
                    d.add_paragraph(f"{k} : {v}")
            else:
                d.add_paragraph(_s(value))

    # --- Legal mentions ---
    if legal := c["legal_mentions"]:
        d.add_paragraph()
        add_h2("Mentions légales")
        p = d.add_paragraph(legal)
        for run in p.runs:
            run.font.size = Pt(8)

    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf.read()


# ===========================================================================
# XLSX exporter — openpyxl
# ===========================================================================


def _export_xlsx(doc: Any, company: Any) -> bytes:
    try:
        import openpyxl
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
    except ImportError as exc:
        raise ExportUnavailable(
            "openpyxl n'est pas installé (pip install openpyxl)."
        ) from exc

    c = _ch(company)
    content: dict = getattr(doc, "content", None) or {}
    doc_type: str = _s(getattr(doc, "document_type", ""))
    title: str = _s(getattr(doc, "title", "Document")) or "Document"

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Devis" if doc_type == "quote" else "Document"

    # Shared style helpers
    AMBER = "D97706"
    LIGHT = "F3F4F6"
    WHITE = "FFFFFF"

    font_bold = Font(bold=True)
    font_hdr = Font(bold=True, color=WHITE)
    fill_amber = PatternFill("solid", fgColor=AMBER)
    fill_grey = PatternFill("solid", fgColor=LIGHT)
    align_r = Alignment(horizontal="right")
    thin = Side(style="thin", color="D1D5DB")
    cell_border = Border(left=thin, right=thin, top=thin, bottom=thin)

    row = 1

    # --- Company header ---
    for line in _company_lines(company):
        ws.cell(row=row, column=1, value=line)
        row += 1
    row += 1  # blank line

    # --- Document title + date ---
    cell = ws.cell(row=row, column=1, value=title)
    cell.font = Font(bold=True, size=12)
    row += 1
    ws.cell(row=row, column=1, value=f"Date : {_today()}")
    row += 2  # blank line before content

    if doc_type == "quote":
        lines, tva_rate, total_ht, total_tva, total_ttc = _parse_quote(content, c["default_tva_rate"])

        # Column headers
        headers = ["Désignation", "Qté", "Unité", "PU HT (€)", "Total HT (€)"]
        col_widths = [42, 10, 10, 18, 18]
        for col_i, (hdr, w) in enumerate(zip(headers, col_widths), start=1):
            cell = ws.cell(row=row, column=col_i, value=hdr)
            cell.font = font_hdr
            cell.fill = fill_amber
            cell.border = cell_border
            cell.alignment = Alignment(horizontal="center")
            ws.column_dimensions[get_column_letter(col_i)].width = w
        row += 1

        # Data rows
        for i, ln in enumerate(lines):
            fill = fill_grey if i % 2 else None
            values = [
                ln["label"],
                ln["qty"],
                ln["unit"],
                ln["unit_price_ht"],
                ln["total_ht"],
            ]
            for col_i, val in enumerate(values, start=1):
                cell = ws.cell(row=row, column=col_i, value=val)
                cell.border = cell_border
                if fill:
                    cell.fill = fill
                if col_i in (2, 4, 5):
                    cell.alignment = align_r
                if col_i in (4, 5) and isinstance(val, float):
                    cell.number_format = '#,##0.00 "€"'
            row += 1

        row += 1  # blank

        # Totals
        tva_pct = f"{tva_rate * 100:.0f} %"
        for label, amount in [
            ("Total HT", total_ht),
            (f"TVA ({tva_pct})", total_tva),
            ("Total TTC", total_ttc),
        ]:
            lbl_cell = ws.cell(row=row, column=4, value=label)
            amt_cell = ws.cell(row=row, column=5, value=round(amount, 2))
            amt_cell.number_format = '#,##0.00 "€"'
            amt_cell.alignment = align_r
            if label.startswith("Total TTC"):
                lbl_cell.font = Font(bold=True)
                amt_cell.font = Font(bold=True)
            row += 1

        # Hypotheses
        hypotheses = _lst(content.get("hypotheses"))
        if hypotheses:
            row += 1
            ws.cell(row=row, column=1, value="Hypothèses").font = font_bold
            row += 1
            for h in hypotheses:
                ws.cell(row=row, column=1, value=f"• {_s(h)}")
                row += 1

    else:
        # Generic / report: section title in col A, values in col B
        ws.column_dimensions["A"].width = 28
        ws.column_dimensions["B"].width = 60

        if doc_type == "site_report":
            sections: list[tuple[str, Any]] = [
                ("Date de visite", content.get("date", "")),
                ("Participants", content.get("present", [])),
                ("Constats", content.get("constats", [])),
                ("Actions à mener", content.get("actions", [])),
                ("Réserves", content.get("reserves", [])),
            ]
        elif doc_type == "tender_response":
            sections = [
                ("Pièces demandées", content.get("pieces_demandees", [])),
                ("Critères de sélection", content.get("criteres", [])),
                ("Délais", content.get("delais", "")),
                ("Points de vigilance", content.get("points_vigilance", [])),
            ]
        elif doc_type == "photo_report":
            sections = [
                ("Observations", content.get("observations", "")),
                ("Travaux visibles", content.get("travaux_visibles", [])),
                ("Points d'attention", content.get("points_attention", [])),
            ]
        else:
            sections = [(k.replace("_", " ").capitalize(), v) for k, v in content.items()]

        for section_label, value in sections:
            if isinstance(value, list):
                if not value:
                    continue
                ws.cell(row=row, column=1, value=section_label).font = font_bold
                for item in value:
                    ws.cell(row=row, column=2, value=f"• {_s(item)}")
                    row += 1
            else:
                text = _s(value)
                if not text:
                    continue
                ws.cell(row=row, column=1, value=section_label).font = font_bold
                ws.cell(row=row, column=2, value=text)
                row += 1

    # Legal mentions (for all types)
    if legal := c["legal_mentions"]:
        row += 1
        ws.cell(row=row, column=1, value="Mentions légales").font = font_bold
        row += 1
        for legal_line in legal.splitlines():
            if legal_line.strip():
                cell = ws.cell(row=row, column=1, value=legal_line)
                cell.font = Font(size=8)
                row += 1

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


# ===========================================================================
# Public entrypoint
# ===========================================================================

_MEDIA: dict[str, str] = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}


def export_document(
    doc: Any,
    company: Any,
    fmt: str,
) -> tuple[bytes, str, str]:
    """Export a Document ORM row to bytes in the requested format.

    Args:
        doc: ``Document`` ORM row (or any object with ``title``,
            ``document_type``, ``content`` attributes).
        company: ``CompanySettings`` ORM row, or ``None``.  When ``None``, a
            sensible default company name is used and no crash occurs.
        fmt: ``"pdf"``, ``"docx"``, or ``"xlsx"`` (case-insensitive).

    Returns:
        ``(bytes, media_type, filename)`` — ready to pass to a FastAPI
        ``StreamingResponse`` with ``Content-Disposition: attachment``.

    Raises:
        ValueError: for an unrecognised *fmt*.
        ExportUnavailable: if the underlying library is not installed.
    """
    fmt = fmt.lower().strip()
    if fmt not in _MEDIA:
        raise ValueError(
            f"Format non supporté : {fmt!r}. Valeurs acceptées : pdf, docx, xlsx."
        )

    title = _s(getattr(doc, "title", "document")) or "document"
    stem = _filename_stem(title)
    filename = f"{stem}.{fmt}"
    media_type = _MEDIA[fmt]

    if fmt == "pdf":
        data = _export_pdf(doc, company)
    elif fmt == "docx":
        data = _export_docx(doc, company)
    else:
        data = _export_xlsx(doc, company)

    return data, media_type, filename
